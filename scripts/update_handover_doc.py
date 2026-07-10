from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


DOC_PATH = Path(r"C:\Users\19430\OneDrive\Documents\NA Intern\国金\实习工作交接文档.docx")
OUT_PATH = Path(r"C:\Users\19430\OneDrive\Documents\NA Intern\国金\实习工作交接文档_项目一已写入.docx")


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph, text: str = "", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph, rows):
    document = paragraph._parent
    table = document.add_table(rows=1, cols=len(rows[0]), width=Inches(6))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.autofit = True
    table._tbl.getparent().remove(table._tbl)
    paragraph._p.addnext(table._tbl)

    first = table.rows[0].cells
    for i, value in enumerate(rows[0]):
        first[i].text = value
        for run in first[i].paragraphs[0].runs:
            run.bold = True

    for row in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    return table


def set_font(paragraph, size=None, bold=None):
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        if size is not None:
            run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold


def add_heading(after, text, level=2):
    style = f"Heading {level}" if level <= 3 else None
    para = insert_paragraph_after(after, text, style=style)
    set_font(para, size=14 if level == 2 else 12, bold=True)
    return para


def add_body(after, text):
    para = insert_paragraph_after(after, text)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    set_font(para, size=10.5)
    return para


def add_bullet(after, text):
    para = insert_paragraph_after(after, f"• {text}")
    para.paragraph_format.left_indent = Pt(14)
    para.paragraph_format.space_after = Pt(3)
    set_font(para, size=10.5)
    return para


def add_code(after, text):
    para = insert_paragraph_after(after, text)
    para.paragraph_format.left_indent = Pt(18)
    para.paragraph_format.space_after = Pt(3)
    run = para.runs[0] if para.runs else para.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)
    return para


def add_small_table(after, rows):
    table = insert_table_after(after, rows)
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                set_font(para, size=9.5)
                para.paragraph_format.space_after = Pt(2)
    return table


def main():
    doc = Document(DOC_PATH)
    paragraphs = doc.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip().startswith("三、交接项目一"))
    end = next(i for i, p in enumerate(paragraphs) if p.text.strip().startswith("四、交接项目二"))

    anchor = paragraphs[start - 1]
    following_xml = deepcopy(paragraphs[end]._p)

    for p in paragraphs[start:end]:
        delete_paragraph(p)

    current = anchor
    current = add_heading(current, "三、交接项目一：Skill 自动评估系统", level=1)

    current = add_heading(current, "3.1 项目背景", level=2)
    for text in [
        "公司在推进人工智能能力落地时，各部门会逐步沉淀不同类型的 Skill，但目前存在比较明显的三类问题：",
        "第一，认知门槛较高。多数员工并不熟悉人工智能工作流、智能体、Skill 等概念，即使已有可用能力，也不一定知道如何使用、如何判断是否适合自己的业务场景。",
        "第二，资产较为分散。不同部门、不同个人会根据各自需求沉淀 Skill，但这些资产可能散落在不同位置，缺少统一整理、归纳和共享方式，后续复用成本较高。",
        "第三，质量缺少统一标准。Skill 的说明质量、输入输出格式、适用边界、测试样例和安全约束差异较大。好用的 Skill 不容易被发现，质量不稳定的 Skill 又可能误导业务使用。",
        "在这个背景下，Skill 自动评估系统的定位是：在 Skill 被共享、复用或未来上架前，先建立一套统一的质量检查和评估流程。系统通过材料检查、风险判断、测试题补齐、双模型评审和专家复核，帮助团队判断一个 Skill 是否达到可复用的质量底线。",
        "本项目重点解决的是“如何评估一个 Skill 是否合格”，而不是直接建设完整的 Skill 集市。当前已完成的是 Skill 评估系统主链路，为后续统一管理和共享 Skill 打基础。",
        "项目 GitHub 地址：https://github.com/NathanZhou1225/Skillhub_eval/tree/main",
        r"本地仓库路径：C:\Users\19430\OneDrive\Documents\NA Intern\国金\Skillhub",
    ]:
        current = add_body(current, text)

    current = add_heading(current, "3.2 项目目标", level=2)
    for text in [
        "建立统一的 Skill 准入标准：明确一个 Skill 在被复用前，应该具备哪些说明、输入输出样例、测试题、安全边界和评估材料。",
        "降低 Skill 作者准备材料的门槛：作者不需要一开始就完全理解全部评估规则，可以通过对话式流程上传压缩包、补充材料、确认补题计划，并由系统辅助补齐测试题。",
        "提升评估结果的一致性和可解释性：系统通过规则检查、题型门槛、两个评审模型交叉评估和专家复核，输出“通过 / 需人工复核 / 不通过”结论，并在报告中展示原因、分数、分歧和风险提示。",
        "支持真实执行验证：除了读取作者提供的样例输出，系统也支持调用本机 Codex、Cursor Agent、Trae 等命令行工具真实执行测试题，再把真实产出交给同一套评估流程判断。",
        "为后续 Skill 共享和集市建设打基础：当前阶段主要完成评估系统。后续可以基于评估结果继续建设服务器部署、Skill 上架、分类浏览、自然语言搜索和上架后健康检查等能力。",
    ]:
        current = add_bullet(current, text)

    current = add_heading(current, "3.3 项目当前完成情况", level=2)
    current = add_body(current, "当前项目整体分为四个阶段：")
    phase_rows = [
        ["阶段", "当前状态", "说明"],
        ["阶段一：规范文档", "已完成", "包括 Skill 元数据定义、编写规范、评估指标、准入标准、评审流程等基础规则。"],
        ["阶段二：评估引擎", "已完成", "系统可以完成结构检查、材料缺口扫描、规则校验、双模型评审、报告生成和专家复核等核心评估能力。"],
        ["阶段三：评估系统完善", "主链路已验收", "已支持对话式上传、材料补充、自动补题、正式评估、报告展示、本地工具真实执行和环境检查等能力。"],
        ["阶段四：服务器部署和 Skill 集市", "尚未启动", "后续可继续推进服务器多人访问、Skill 上架、分类浏览、自然语言搜索、上架后健康检查等能力。"],
    ]
    table = add_small_table(current, phase_rows)
    current = table._tbl
    # Move insertion cursor back to the paragraph following table by inserting after table XML manually.
    temp_para = doc.add_paragraph()
    temp_para._p.getparent().remove(temp_para._p)
    table._tbl.addnext(temp_para._p)
    current_para = temp_para

    current = add_body(current_para, "当前可演示主流程：作者上传 Skill 压缩包 → 系统检查包结构、材料完整性和安全风险 → 如果测试题不足，系统展示补题计划 → 作者确认后，系统自动补题或引导对话补充 → 条件满足后自动进入正式评估 → 两个评审模型分别评分 → 系统汇总结果，生成“通过 / 需人工复核 / 不通过”结论 → 如果结论为“需人工复核”，可由专家视角进行批准或退回。")
    current = add_body(current, "当前已完成的主要能力包括：")
    for text in [
        "Skill 压缩包上传与对话式评估入口",
        "材料完整性检查、安全扫描和风险等级判断",
        "测试题型门槛检查、自动补题计划和作者对话补充材料",
        "正式双模型评估，以及“通过 / 需人工复核 / 不通过”结论生成",
        "报告详情、评分过程、用量与耗时展示",
        "专家视角复核和会话归档",
        "本地命令行工具真实执行能力",
        "本地执行环境检查，作为诊断工具使用",
        "启动说明、运行手册、全景说明和项目记录文档整理",
    ]:
        current = add_bullet(current, text)
    current = add_body(current, "尚未完成或后续可继续推进：")
    for text in [
        "阶段三正式收官确认",
        "服务器部署和多人访问",
        "Skill 集市页面、上架发布流程和使用侧自然语言搜索",
        "上架后健康检查",
        "根据真实业务反馈继续优化评估体验和报告呈现",
        "补充更多真实业务样例，验证评估机制是否适合实际工作流",
    ]:
        current = add_bullet(current, text)

    current = add_heading(current, "3.4 项目使用方式", level=2)
    for text in [
        "当前项目代码托管在 GitHub，接手人应从仓库重新拉取代码并在自己的电脑或服务器上配置运行环境。",
        "项目 GitHub 地址：https://github.com/NathanZhou1225/Skillhub_eval/tree/main",
        "本地运行时建议参考仓库根目录的 README.md。核心步骤如下：",
    ]:
        current = add_body(current, text)
    for text in [
        "拉取 GitHub 仓库代码。",
        "确认本机已安装 Python 3.11 或以上版本。建议优先使用 Python 3.11 或 3.12。",
        "在项目根目录安装项目依赖和命令入口：",
    ]:
        current = add_bullet(current, text)
    current = add_code(current, 'pip install -e ".[dev]"')
    current = add_body(current, "说明：项目依赖已经写在 pyproject.toml 中，目前不需要额外维护 requirements.txt。如果后续公司部署流程强制要求 requirements.txt，再从 pyproject.toml 单独整理即可。")
    current = add_bullet(current, "从 .env.example 复制生成 .env，并配置评审模型相关参数：")
    current = add_code(current, "copy .env.example .env")
    current = add_bullet(current, "检查评审模型是否可连接：")
    current = add_code(current, "python scripts/check_providers.py")
    current = add_bullet(current, "启动服务。推荐写法：")
    current = add_code(current, "python -m skillhub_eval.adapters.cli.main serve --host 127.0.0.1 --port 8000")
    current = add_body(current, "如果已经通过 pip install -e \".[dev]\" 正确安装项目，也可以使用：")
    current = add_code(current, "skillhub-eval serve")
    current = add_body(current, "如果 PowerShell 提示 skillhub-eval 无法识别，通常说明当前环境还没有安装该项目，或没有进入正确的 Python 环境。此时应先重新执行安装命令，或使用上面的 python -m 启动方式。")
    current = add_bullet(current, "打开浏览器访问评估页面：")
    current = add_code(current, "http://127.0.0.1:8000/ui/index.html")
    current = add_bullet(current, "接口文档页面：")
    current = add_code(current, "http://127.0.0.1:8000/docs")
    current = add_body(current, "如果需要让局域网内其他电脑访问，应将服务绑定到 0.0.0.0：")
    current = add_code(current, "python -m skillhub_eval.adapters.cli.main serve --host 0.0.0.0 --port 8000")
    current = add_body(current, "其他电脑访问：")
    current = add_code(current, "http://运行服务机器的IP地址:8000/ui/index.html")
    current = add_body(current, "使用时需要注意：")
    for text in [
        "127.0.0.1 只代表运行浏览器的当前电脑。如果服务在 A 电脑运行，B 电脑不能用 127.0.0.1 访问 A 电脑服务。",
        "换电脑或服务器承接时，需要重新配置 .env、模型密钥、本地路径、数据库路径和本地命令行工具登录状态。",
        "本地真实执行是可选增强；没有配置 Codex、Cursor Agent、Trae 等本地工具时，也可以使用样例输出完成评估。",
        "页面中的“环境检查”只是诊断工具，不是正式评估门禁。",
        "修改 .env 后需要重启服务。",
    ]:
        current = add_bullet(current, text)

    current = add_heading(current, "3.5 仓库和资料说明", level=2)
    current = add_body(current, "项目仓库主要分为“运行说明类文档、产品说明类文档、规范类文档、过程记录类文档、功能代码、测试样例”几类。")

    sections = [
        ("3.5.1 根目录重点文件", [
            ("README.md", "接手人首先阅读的启动说明。包含项目定位、环境要求、快速开始、新机器启动检查清单、换电脑或服务器承接注意事项、项目结构和常用命令。"),
            ("RECORD.md", "项目总账和交接快照。用于记录项目目标、阶段状态、重要决策、已完成事项、未完成事项和历史推进记录。接手人应优先阅读顶部“交接快照”，不要直接把历史已完成清单当成当前待办。"),
            (".env.example", "环境变量模板。用于配置评审模型、数据库路径、评估材料暂存路径、本地执行偏好等。正式 .env 不应提交到仓库。"),
            ("pyproject.toml", "项目依赖、测试和打包配置。偏技术配置，产品接手人了解其存在即可。"),
        ]),
        ("3.5.2 docs/：产品说明、规范和运行文档", [
            ("docs/Project-Background.md", "项目业务背景说明。解释公司推进人工智能能力落地时面临的认知门槛、资产碎片化和质量无标准等问题。"),
            ("docs/guides/Skill评估系统全景说明.md", "最适合产品、业务和接手人阅读的系统全貌说明。解释系统为什么做、完整流程是什么、如何评估、哪些能力已经具备、哪些边界不能保证。"),
            ("docs/guides/Skill评估机制架构与流程.md", "说明评估系统的整体流程和架构逻辑，适合理解从上传到出结论的主链路。"),
            ("docs/guides/Skills评估标准白皮书.md", "偏标准解释和对外说明，适合用于汇报或解释为什么这样评估 Skill。"),
            ("docs/guides/本地CLI Agent真跑机制说明.md", "说明为什么要支持本地 Codex、Cursor Agent、Trae 等工具真实执行，以及样例评估和真实执行评估的区别。"),
            ("docs/guides/报告呈现规范.md", "说明评估报告中各类卡片、原因、分歧、耗时和展示口径。"),
            ("docs/guides/Skill编写指南.md", "面向 Skill 作者，说明如何准备 Skill 材料、如何写说明、如何处理退回和补题。"),
            ("docs/specs/Skill元数据定义与编写规范.md", "Skill 包结构和字段规范。用于定义一个合格 Skill 包应该包含什么。"),
            ("docs/specs/评估指标与准入标准.md", "评分公式、通过阈值和红线规则的权威文档。后续如果讨论“为什么通过 / 为什么不通过”，应优先查这里。"),
            ("docs/specs/评审Agent工作流与Prompt骨架.md", "说明评审流程和评审提示词结构，偏机制设计。"),
            ("docs/runbooks/local-agent-exec-validation.md", "本地命令行工具真实执行能力的验收说明。用于验证 Codex、Cursor Agent、Trae 等工具是否能真实执行测试题。"),
            ("docs/research/Skill数据定义与编写规范调研.md", "前期调研资料，说明 Skill 数据定义和编写规范的参考依据。"),
        ]),
        ("3.5.3 .project_memory/：项目记忆和阶段记录", [
            (".project_memory/active/SPRINT_phase3-eval-system.md", "阶段三评估系统的任务真源。记录从对话评估、本地演示、本地真实执行到运行环境产品化的推进情况。"),
            (".project_memory/active/SPRINT_phase4-marketplace-biz.md", "阶段四计划。当前尚未启动，主要包括服务器部署、Skill 集市、上架、自然语言搜索和上架后健康检查。"),
            (".project_memory/backlog/BACKLOG.md", "非当前主线的后续优化项。里面的内容不等于当前必须立即做的待办。"),
            (".project_memory/global/ARCHITECTURE.md", "架构说明，偏技术，但可以帮助接手人理解模块边界。"),
            (".project_memory/archive/", "已归档阶段记录，用于历史追溯。"),
        ]),
        ("3.5.4 openspec/：变更记录和规格沉淀", [
            ("openspec/specs/skill-execution/spec.md", "当前本地执行能力的主规格文档。说明本地工具真实执行、失败处理、报告归属等机制。"),
            ("openspec/changes/archive/", "已完成并归档的功能变更。每个归档目录通常包含 proposal.md、design.md、tasks.md 和规格变更，用于追溯某个功能为什么这样设计。"),
        ]),
        ("3.5.5 skillhub_eval/：系统功能代码", [
            ("core/", "评估主逻辑，包括材料检查、风险判断、补题、评分、报告生成等。"),
            ("adapters/", "系统入口，包括网页接口、命令行入口和页面资源。"),
            ("providers/", "评审模型接入，例如 DeepSeek、Gemini 或兼容 OpenAI 调用方式的模型接口。"),
            ("execution/", "本地命令行工具真实执行相关能力，包括 Codex、Cursor Agent、Trae 等本地工具的检测、选择、执行和结果归因。"),
            ("persistence/", "数据记录能力，主要记录评估轮次、会话、模型投票、专家操作等。"),
            ("settings.py", "环境变量和配置读取。"),
        ]),
        ("3.5.6 tests/ 和 testskills/：测试与样例", [
            ("tests/", "自动化测试目录，用于验证系统功能是否正常。偏技术维护，不建议产品接手人逐项阅读。"),
            ("testskills/", "样例 Skill 和测试包，用于本地演示或回归验证。包括 stock-radar、grill-me、tiered-memory-sprint-manager 等样例。"),
        ]),
    ]
    for heading, items in sections:
        current = add_heading(current, heading, level=3)
        for name, desc in items:
            current = add_bullet(current, f"{name}：{desc}")

    current = add_heading(current, "3.6 交接注意事项", level=2)
    for text in [
        "当前交付物是“Skill 自动评估系统”，不是完整 Skill 集市平台。阶段四中的服务器部署、Skill 集市、Skill 上架和消费者搜索尚未启动。",
        "阶段三主链路已经可演示，但仍可继续增强。当前主链路已实机验收，能够完成上传、补题、正式评估、报告和专家复核。本地真实执行也已经作为可选路径跑通。",
        "本地真实执行是可选增强，不是默认硬要求。系统默认可以使用 Skill 包中的样例输出进行评估。如果选择本地真实执行，则需要运行服务的电脑安装并登录对应命令行工具。",
        "环境检查只是诊断，不是正式评估门禁。页面中的“运行环境检查”用于提前发现本机工具、模型或基础执行问题。检查失败不代表 Skill 一定不能评估，正式结论仍以评估流程为准。",
        "换电脑或服务器时要重点看 README。新机器接手时，应重新配置 .env、模型密钥、本地路径、数据库路径和本地工具登录状态。不要直接复用旧电脑的绝对路径或运行缓存。",
        "RECORD.md 顶部交接快照最重要。接手人不要从历史流水中找当前状态，容易误把已完成或已归档的问题当成待办。应先看 RECORD.md 顶部交接快照，再看 Sprint。",
        "历史测试数量是阶段快照，不是当前基线。文档中出现的 235 tests、400 tests、742 passed 等是当时里程碑记录，不应简单相加，也不一定代表当前全量测试数。",
        "不建议随意修改评分阈值。评估指标与准入标准已经锁定，包括通过线、红线题和双模型分歧规则。后续如需调整，应先形成明确产品决策。",
    ]:
        current = add_bullet(current, text)

    current = add_heading(current, "3.7 后续优化建议", level=2)
    for text in [
        "阶段三正式收官：在现有主链路已验收的基础上，整理最后的验收口径，确认是否正式关闭阶段三。",
        "服务器部署：将当前本地服务部署到服务器，支持多人通过浏览器访问。需要重点处理服务地址、端口、防火墙、模型密钥、数据库路径和本地执行边界。",
        "Skill 集市：在评估通过后，增加 Skill 列表、上架、分类浏览、搜索、热度展示等能力，让通过评估的 Skill 可以被业务用户发现和使用。",
        "上架后健康检查：对已通过的 Skill 固化标准测试样例，后续定期重跑，防止模型、接口或业务环境变化导致 Skill 失效。",
        "产品化报告优化：继续优化报告的业务表达，使非技术用户能更快看懂为什么通过、为什么需人工复核、为什么不通过。",
        "真实业务样例补充：当前已有测试样例，但后续若要推广，应引入更多真实业务场景，用于验证评估机制是否符合实际工作流。",
    ]:
        current = add_bullet(current, text)

    current = add_heading(current, "3.8 建议接手人阅读顺序", level=2)
    reading = [
        "README.md：先了解项目是什么、怎么启动、怎么换电脑承接。",
        "RECORD.md 顶部“交接快照”：了解当前真实状态、已完成内容、未完成内容和下一步优先级。",
        "docs/Project-Background.md：了解项目业务背景和整体痛点。",
        "docs/guides/Skill评估系统全景说明.md：从产品视角理解系统为什么做、怎么评估、当前能力边界是什么。",
        ".project_memory/active/SPRINT_phase3-eval-system.md：查看阶段三评估系统的完整推进记录。",
        ".project_memory/active/SPRINT_phase4-marketplace-biz.md：如果准备继续做服务器或集市，再阅读阶段四计划。",
        "docs/specs/评估指标与准入标准.md：如果需要解释评分、阈值、红线或人工复核规则，再阅读该文档。",
        "docs/runbooks/local-agent-exec-validation.md：如果需要验证本地 Codex、Cursor Agent、Trae 真实执行能力，再阅读该文档。",
    ]
    for idx, text in enumerate(reading, start=1):
        current = add_body(current, f"{idx}. {text}")

    current._p.addnext(following_xml)
    try:
        doc.save(DOC_PATH)
    except PermissionError:
        doc.save(OUT_PATH)


if __name__ == "__main__":
    main()
